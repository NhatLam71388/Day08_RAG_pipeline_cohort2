"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.

Chiến lược: Tạo 3 file DOCX chứa nội dung tóm tắt các văn bản pháp luật
thực tế về phòng chống ma tuý tại Việt Nam.

Văn bản được tổng hợp từ:
    - Luật Phòng, chống ma tuý 2021 (73/2021/QH15)
    - Nghị định 105/2021/NĐ-CP
    - Bộ luật Hình sự 2015 (sửa đổi 2017) - Chương XX
"""

from pathlib import Path

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")


def create_legal_docs():
    """Tạo 3 file DOCX chứa nội dung văn bản pháp luật về ma tuý."""
    try:
        from docx import Document
    except ImportError:
        print("⚠ python-docx chưa được cài. Chạy: pip install python-docx")
        return

    docs = [
        {
            "filename": "luat-phong-chong-ma-tuy-2021.docx",
            "title": "LUẬT PHÒNG, CHỐNG MA TUÝ 2021 (Luật số 73/2021/QH15)",
            "content": [
                ("Chương I: QUY ĐỊNH CHUNG", [
                    "Điều 1. Phạm vi điều chỉnh",
                    "Luật này quy định về phòng ngừa, ngăn chặn, đấu tranh chống tệ nạn ma tuý; kiểm soát các hoạt động hợp pháp liên quan đến ma tuý; cai nghiện ma tuý; trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức và Nhà nước trong phòng, chống ma tuý.",
                    "Điều 2. Giải thích từ ngữ",
                    "1. Chất ma tuý là chất gây nghiện, chất hướng thần được quy định trong danh mục chất ma tuý do Chính phủ ban hành.",
                    "2. Chất gây nghiện là chất kích thích hoặc ức chế thần kinh, dễ gây tình trạng nghiện đối với người sử dụng.",
                    "3. Tiền chất là hóa chất không thể thiếu trong quá trình điều chế, sản xuất chất ma tuý.",
                    "Điều 3. Nguyên tắc phòng, chống ma tuý",
                    "1. Lấy phòng ngừa là chính; chú trọng tuyên truyền, giáo dục pháp luật và vận động nhân dân tham gia phòng, chống ma tuý.",
                    "2. Cai nghiện ma tuý là biện pháp quan trọng trong phòng, chống ma tuý.",
                    "Điều 4. Chính sách của Nhà nước về phòng, chống ma tuý",
                    "Nhà nước có chính sách huy động mọi lực lượng trong xã hội tham gia vào công tác phòng, chống ma tuý.",
                ]),
                ("Chương II: PHÒNG NGỪA TỆ NẠN MA TUÝ", [
                    "Điều 10. Trách nhiệm của Mặt trận Tổ quốc Việt Nam",
                    "Mặt trận Tổ quốc Việt Nam trong phạm vi nhiệm vụ, quyền hạn của mình có trách nhiệm tuyên truyền, vận động nhân dân thực hiện pháp luật về phòng, chống ma tuý.",
                    "Điều 11. Trách nhiệm của gia đình",
                    "1. Giáo dục thành viên trong gia đình về tác hại của ma tuý và cách phòng, chống tệ nạn ma tuý.",
                    "2. Quản lý, giám sát thành viên trong gia đình, phòng ngừa thành viên trong gia đình vi phạm pháp luật về phòng, chống ma tuý.",
                ]),
                ("Chương III: CỨU GIÚP NGƯỜI NGHIỆN MA TUÝ", [
                    "Điều 28. Cai nghiện ma tuý tự nguyện",
                    "1. Người nghiện ma tuý tự nguyện đăng ký cai nghiện tại cơ sở cai nghiện ma tuý.",
                    "2. Thời gian cai nghiện ma tuý tự nguyện tối thiểu là 06 tháng.",
                    "Điều 29. Cai nghiện ma tuý bắt buộc",
                    "1. Người nghiện ma tuý từ đủ 18 tuổi trở lên bị áp dụng biện pháp xử lý hành chính đưa vào cơ sở cai nghiện bắt buộc.",
                    "2. Thời gian cai nghiện bắt buộc từ 12 tháng đến 24 tháng.",
                    "Điều 30. Quản lý sau cai nghiện",
                    "Người sau cai nghiện ma tuý được quản lý tại nơi cư trú trong thời gian tối thiểu 12 tháng kể từ ngày chấp hành xong quyết định cai nghiện bắt buộc.",
                ]),
            ]
        },
        {
            "filename": "nghi-dinh-105-2021-nd-cp.docx",
            "title": "NGHỊ ĐỊNH 105/2021/NĐ-CP\nQuy định chi tiết và hướng dẫn thi hành một số điều của Luật Phòng, chống ma tuý",
            "content": [
                ("Chương I: QUY ĐỊNH CHUNG", [
                    "Điều 1. Phạm vi điều chỉnh",
                    "Nghị định này quy định chi tiết và hướng dẫn thi hành một số điều của Luật Phòng, chống ma tuý năm 2021 về: quản lý người sử dụng trái phép chất ma tuý; cai nghiện ma tuý tự nguyện tại gia đình và cộng đồng; cai nghiện ma tuý tự nguyện tại cơ sở cai nghiện ma tuý.",
                    "Điều 2. Đối tượng áp dụng",
                    "1. Người sử dụng trái phép chất ma tuý.",
                    "2. Cơ quan, tổ chức, cá nhân có liên quan đến công tác phòng, chống ma tuý.",
                ]),
                ("Chương II: QUẢN LÝ NGƯỜI SỬ DỤNG TRÁI PHÉP CHẤT MA TUÝ", [
                    "Điều 5. Xác định tình trạng nghiện ma tuý",
                    "1. Việc xác định tình trạng nghiện ma tuý được thực hiện theo hướng dẫn của Bộ Y tế.",
                    "2. Người nghiện ma tuý là người sử dụng chất ma tuý, thuốc gây nghiện, thuốc hướng thần và bị lệ thuộc vào các chất này.",
                    "Điều 8. Quản lý người sử dụng trái phép chất ma tuý tại cộng đồng",
                    "1. Người sử dụng trái phép chất ma tuý được giao cho gia đình quản lý trong thời gian xem xét quyết định áp dụng biện pháp cai nghiện.",
                    "2. Thời gian quản lý tại gia đình không quá 30 ngày.",
                    "Điều 15. Hỗ trợ chi phí cai nghiện",
                    "Ngân sách nhà nước hỗ trợ chi phí cai nghiện ma tuý đối với người không có khả năng tài chính.",
                ]),
                ("Chương III: CAI NGHIỆN MA TUÝ TỰ NGUYỆN TẠI GIA ĐÌNH VÀ CỘNG ĐỒNG", [
                    "Điều 20. Điều kiện cai nghiện tại gia đình",
                    "1. Người cai nghiện tại gia đình phải có đủ điều kiện về chỗ ở và người quản lý.",
                    "2. Thời gian cai nghiện tại gia đình tối thiểu là 06 tháng.",
                    "Điều 21. Quy trình cai nghiện tại gia đình",
                    "Bước 1: Tiếp nhận và đánh giá tình trạng nghiện.",
                    "Bước 2: Lập kế hoạch cai nghiện cá nhân.",
                    "Bước 3: Thực hiện các biện pháp điều trị và hỗ trợ.",
                    "Bước 4: Đánh giá kết quả và lập kế hoạch tái hoà nhập cộng đồng.",
                ]),
            ]
        },
        {
            "filename": "bo-luat-hinh-su-2015-chuong-xx-ma-tuy.docx",
            "title": "BỘ LUẬT HÌNH SỰ 2015 (SỬA ĐỔI 2017)\nChương XX: Các tội phạm về ma tuý (Điều 247–259)",
            "content": [
                ("CÁC TỘI PHẠM VỀ MA TUÝ", [
                    "Điều 247. Tội trồng cây thuốc phiện, cây côca, cây cần sa hoặc các loại cây khác có chứa chất ma tuý",
                    "1. Người nào trồng cây thuốc phiện, cây côca, cây cần sa hoặc các loại cây khác có chứa chất ma tuý, đã được giáo dục nhiều lần, đã được tạo điều kiện để ổn định cuộc sống và đã bị xử phạt vi phạm hành chính về hành vi này hoặc đã bị kết án về tội này, chưa được xoá án tích mà còn vi phạm, thì bị phạt tù từ 06 tháng đến 03 năm.",
                    "2. Phạm tội trong trường hợp có tổ chức hoặc với quy mô lớn, thì bị phạt tù từ 03 năm đến 07 năm.",
                    "Điều 248. Tội sản xuất trái phép chất ma tuý",
                    "1. Người nào sản xuất trái phép chất ma tuý, thì bị phạt tù từ 02 năm đến 07 năm.",
                    "2. Phạm tội trong các trường hợp: có tổ chức; chất ma tuý có khối lượng lớn; hoặc gây hậu quả nghiêm trọng, thì bị phạt tù từ 07 năm đến 15 năm.",
                    "3. Phạm tội trong trường hợp chất ma tuý có khối lượng rất lớn hoặc gây hậu quả rất nghiêm trọng, thì bị phạt tù từ 15 năm đến 20 năm.",
                    "4. Phạm tội trong trường hợp đặc biệt nghiêm trọng, thì bị phạt tù 20 năm, tù chung thân hoặc tử hình.",
                    "Điều 249. Tội tàng trữ trái phép chất ma tuý",
                    "1. Người nào tàng trữ trái phép chất ma tuý mà không nhằm mục đích mua bán, vận chuyển, thì bị phạt tù từ 01 năm đến 05 năm.",
                    "2. Phạm tội trong trường hợp: có tổ chức; chất ma tuý khối lượng lớn; hoặc tái phạm nguy hiểm, thì bị phạt tù từ 05 năm đến 10 năm.",
                    "3. Phạm tội trong trường hợp chất ma tuý khối lượng rất lớn, thì bị phạt tù từ 10 năm đến 15 năm.",
                    "4. Phạm tội trong trường hợp đặc biệt nghiêm trọng, thì bị phạt tù từ 15 năm đến 20 năm.",
                    "Điều 250. Tội vận chuyển trái phép chất ma tuý",
                    "1. Người nào vận chuyển trái phép chất ma tuý, thì bị phạt tù từ 02 năm đến 07 năm.",
                    "2. Phạm tội trong các trường hợp tăng nặng: phạt tù từ 07 năm đến 15 năm.",
                    "3. Phạm tội rất nghiêm trọng: phạt tù từ 15 năm đến 20 năm.",
                    "4. Phạm tội đặc biệt nghiêm trọng: phạt tù 20 năm, chung thân hoặc tử hình.",
                    "Điều 251. Tội mua bán trái phép chất ma tuý",
                    "1. Người nào mua bán trái phép chất ma tuý, thì bị phạt tù từ 02 năm đến 07 năm.",
                    "2. Phạm tội có tổ chức hoặc tái phạm nguy hiểm: phạt tù từ 07 đến 15 năm.",
                    "3. Phạm tội đặc biệt nghiêm trọng: tử hình.",
                    "Điều 255. Tội sử dụng trái phép chất ma tuý",
                    "1. Người nào sử dụng trái phép chất ma tuý dưới bất kỳ hình thức nào, đã bị xử phạt vi phạm hành chính về hành vi này hoặc đã bị kết án về tội này, chưa được xoá án tích mà còn vi phạm, thì bị phạt tù từ 03 tháng đến 02 năm.",
                    "2. Phạm tội từ 02 lần trở lên hoặc tái phạm: phạt tù từ 02 năm đến 05 năm.",
                    "Điều 258. Tội cưỡng bức, lôi kéo người khác sử dụng trái phép chất ma tuý",
                    "1. Người nào cưỡng bức người khác sử dụng trái phép chất ma tuý hoặc lôi kéo người khác sử dụng trái phép chất ma tuý, thì bị phạt tù từ 02 năm đến 07 năm.",
                    "2. Phạm tội đối với người từ đủ 13 tuổi đến dưới 18 tuổi, hoặc phạm tội 02 lần trở lên, hoặc có tổ chức: phạt tù từ 07 năm đến 12 năm.",
                    "3. Phạm tội đối với trẻ em dưới 13 tuổi hoặc gây hậu quả đặc biệt nghiêm trọng: phạt tù từ 12 năm đến 20 năm.",
                ]),
            ]
        },
    ]

    for doc_info in docs:
        filepath = DATA_DIR / doc_info["filename"]
        if filepath.exists():
            print(f"  ↳ Đã tồn tại: {filepath.name}")
            continue

        doc = Document()
        doc.add_heading(doc_info["title"], level=0)

        for section_title, paragraphs in doc_info["content"]:
            doc.add_heading(section_title, level=1)
            for para in paragraphs:
                if para.startswith("Điều"):
                    p = doc.add_paragraph()
                    run = p.add_run(para)
                    run.bold = True
                else:
                    doc.add_paragraph(para)

        doc.save(str(filepath))
        print(f"  ✓ Đã tạo: {filepath.name} ({filepath.stat().st_size:,} bytes)")


if __name__ == "__main__":
    print("=" * 60)
    print("Task 1: Thu thập văn bản pháp luật về ma tuý")
    print("=" * 60)
    setup_directory()
    create_legal_docs()
    print("\n✓ Hoàn thành Task 1!")
